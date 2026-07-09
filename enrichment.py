"""
enrichment.py
─────────────
Module complémentaire pour le SEO Brief Generator.

Ajoute deux fonctionnalités indépendantes des modules existants :

1. Mots-clés secondaires liés au mot-clé principal
   -> generate_secondary_keywords()

2. Maillage interne (recommandations de liens internes)
   - "ancre + lien" quand une page pertinente existe déjà sur le site cible
   - "ancre seule" quand aucun lien n'existe encore (contenu à créer)
   -> fetch_site_urls() + generate_internal_linking()

Le module est volontairement autonome : il n'importe rien des autres
modules du projet (brief_generator, export_utils, site_analyzer...).
Il ne dépend que de : anthropic, requests, python-docx (stdlib pour le reste).
"""

import json
import re
import html as _html
from urllib.parse import urljoin, urlparse
import xml.etree.ElementTree as ET

import requests

# Modèle par défaut pour l'enrichissement (stable, rapide, économique).
# Modifiable si besoin — voir la doc Anthropic pour les identifiants à jour.
DEFAULT_MODEL = "claude-sonnet-4-6"

_LANG_NAMES = {
    "fr": "français", "en": "anglais", "es": "espagnol", "de": "allemand",
    "it": "italien", "pt": "portugais", "nl": "néerlandais", "pl": "polonais",
    "ru": "russe", "ar": "arabe", "ja": "japonais", "zh": "chinois",
}


# ══════════════════════════════════════════════════════════════════════════════
# Helpers internes
# ══════════════════════════════════════════════════════════════════════════════
def _lang_label(language_code: str) -> str:
    return _LANG_NAMES.get((language_code or "fr").lower(), "français")


def _structure_to_outline(structure) -> str:
    """Transforme la structure du brief (liste de dicts) en plan textuel compact."""
    if not structure:
        return "(aucune structure fournie)"
    lines = []
    for item in structure:
        level = item.get("level", "H2")
        title = item.get("title", "")
        try:
            indent = "  " * max(0, int(level[1]) - 2)
        except (ValueError, IndexError):
            indent = ""
        lines.append(f"{indent}{level}: {title}")
    return "\n".join(lines)


def _call_claude(prompt: str, anthropic_key: str, model: str, max_tokens: int = 1500) -> str:
    """Appel simple à l'API Anthropic. Retourne le texte concaténé."""
    import anthropic
    client = anthropic.Anthropic(api_key=anthropic_key)
    resp = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    parts = []
    for block in resp.content:
        # SDK renvoie des blocs typés ; on ne garde que le texte.
        if getattr(block, "type", None) == "text":
            parts.append(block.text)
    return "".join(parts).strip()


def _parse_json(text: str):
    """Parse du JSON même si le modèle a ajouté des ```json ... ``` ou du texte autour."""
    if not text:
        return None
    cleaned = text.strip()
    # Retire les fences Markdown éventuels
    cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # Fallback : extraire le premier objet/tableau JSON de la chaîne
    match = re.search(r"(\{.*\}|\[.*\])", cleaned, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            return None
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 1) MOTS-CLÉS SECONDAIRES
# ══════════════════════════════════════════════════════════════════════════════
def generate_secondary_keywords(
    keyword: str,
    structure=None,
    anthropic_key: str = "",
    language_code: str = "fr",
    max_clusters: int = 5,
    model: str = DEFAULT_MODEL,
) -> dict:
    """
    Génère des mots-clés secondaires liés au mot-clé principal, regroupés par
    thème sémantique, avec pour chaque groupe la section H2/H3 où les intégrer.

    Retour :
    {
      "clusters": [
        {
          "theme": "...",
          "placement": "H2: ...",           # section conseillée
          "keywords": ["...", "...", ...]     # variantes / longue traîne / entités
        },
        ...
      ]
    }
    En cas d'erreur : {"error": "..."}.
    """
    if not keyword or not anthropic_key:
        return {"clusters": []}

    lang = _lang_label(language_code)
    outline = _structure_to_outline(structure)

    prompt = f"""Tu es un expert SEO sémantique. Mot-clé principal : "{keyword}".

Voici le plan (structure Hn) prévu pour l'article :
{outline}

Génère des mots-clés SECONDAIRES en {lang}, directement liés au mot-clé principal,
utiles pour couvrir le champ lexical et l'intention de recherche. Inclus des
variantes, de la longue traîne, des questions et des entités/termes connexes.

Regroupe-les en {max_clusters} clusters thématiques maximum. Pour chaque cluster,
indique la section (H2/H3 du plan ci-dessus) la plus pertinente où intégrer ces
mots-clés. Si aucune section ne convient, propose un intitulé de section à créer.

Réponds UNIQUEMENT avec un objet JSON valide, sans texte ni balises Markdown, au format :
{{
  "clusters": [
    {{
      "theme": "nom court du thème",
      "placement": "H2: intitulé de la section conseillée",
      "keywords": ["mot-clé 1", "mot-clé 2", "mot-clé 3", "mot-clé 4"]
    }}
  ]
}}
Vise 4 à 8 mots-clés par cluster. N'invente pas de chiffres de volume."""

    try:
        raw = _call_claude(prompt, anthropic_key, model, max_tokens=1500)
    except Exception as e:  # noqa: BLE001
        return {"error": f"Erreur API (mots-clés secondaires) : {e}", "clusters": []}

    data = _parse_json(raw)
    if not isinstance(data, dict) or "clusters" not in data:
        # Tolérance : le modèle a peut-être renvoyé directement une liste
        if isinstance(data, list):
            data = {"clusters": data}
        else:
            return {"error": "Réponse non exploitable pour les mots-clés secondaires.",
                    "clusters": []}

    # Nettoyage défensif
    clusters = []
    for c in data.get("clusters", []):
        if not isinstance(c, dict):
            continue
        kws = [str(k).strip() for k in c.get("keywords", []) if str(k).strip()]
        if not kws:
            continue
        clusters.append({
            "theme": str(c.get("theme", "")).strip() or "Mots-clés associés",
            "placement": str(c.get("placement", "")).strip(),
            "keywords": kws,
        })
    return {"clusters": clusters}


# ══════════════════════════════════════════════════════════════════════════════
# 2) MAILLAGE INTERNE
# ══════════════════════════════════════════════════════════════════════════════
def _slug_to_label(url: str) -> str:
    """Déduit un libellé lisible à partir du dernier segment d'une URL."""
    path = urlparse(url).path.rstrip("/")
    slug = path.split("/")[-1] if path else ""
    slug = re.sub(r"\.(html?|php|aspx?)$", "", slug, flags=re.I)
    slug = re.sub(r"[-_]+", " ", slug).strip()
    return slug or url


def fetch_site_urls(site_url: str, max_urls: int = 200, timeout: int = 10) -> list:
    """
    Récupère la liste des pages existantes d'un site via son/ses sitemap(s).

    Tente, dans l'ordre : robots.txt (ligne Sitemap:), /sitemap.xml,
    /sitemap_index.xml. Suit les sitemaps imbriqués (index) sur un niveau.

    Retour : liste de dicts [{"url": "...", "label": "..."}], plafonnée à max_urls.
    Retourne [] si rien n'est trouvé (le maillage se fera alors en "ancre seule").
    """
    if not site_url:
        return []

    parsed = urlparse(site_url if "://" in site_url else "https://" + site_url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    headers = {"User-Agent": "Mozilla/5.0 (compatible; SEOBriefBot/1.0)"}

    candidate_sitemaps = []

    # 1) robots.txt
    try:
        r = requests.get(urljoin(base, "/robots.txt"), headers=headers, timeout=timeout)
        if r.ok:
            for line in r.text.splitlines():
                if line.lower().startswith("sitemap:"):
                    candidate_sitemaps.append(line.split(":", 1)[1].strip())
    except requests.RequestException:
        pass

    # 2) Emplacements classiques
    for path in ("/sitemap.xml", "/sitemap_index.xml", "/sitemap-index.xml"):
        candidate_sitemaps.append(urljoin(base, path))

    seen_sitemaps = set()
    urls = []

    def _parse_sitemap(sm_url: str, depth: int = 0):
        if sm_url in seen_sitemaps or len(urls) >= max_urls or depth > 1:
            return
        seen_sitemaps.add(sm_url)
        try:
            r = requests.get(sm_url, headers=headers, timeout=timeout)
            if not r.ok or not r.content:
                return
            root = ET.fromstring(r.content)
        except (requests.RequestException, ET.ParseError):
            return

        tag = root.tag.lower()
        # Sitemap index -> contient d'autres sitemaps
        if tag.endswith("sitemapindex"):
            for sm in root:
                loc = next((child.text for child in sm
                            if child.tag.lower().endswith("loc") and child.text), None)
                if loc:
                    _parse_sitemap(loc.strip(), depth + 1)
        # Urlset -> contient les pages
        else:
            for url_el in root:
                loc = next((child.text for child in url_el
                            if child.tag.lower().endswith("loc") and child.text), None)
                if loc:
                    loc = loc.strip()
                    urls.append({"url": loc, "label": _slug_to_label(loc)})
                    if len(urls) >= max_urls:
                        return

    for sm in candidate_sitemaps:
        if len(urls) >= max_urls:
            break
        _parse_sitemap(sm)

    # Dédoublonnage en préservant l'ordre
    dedup, seen = [], set()
    for item in urls:
        if item["url"] not in seen:
            seen.add(item["url"])
            dedup.append(item)
    return dedup[:max_urls]


def generate_internal_linking(
    keyword: str,
    structure=None,
    site_pages=None,
    anthropic_key: str = "",
    language_code: str = "fr",
    max_links: int = 6,
    model: str = DEFAULT_MODEL,
) -> dict:
    """
    Génère des recommandations de maillage interne pour le nouvel article.

    - Si une page existante du site cible est pertinente -> "ancre + lien".
    - Sinon (aucun lien encore disponible) -> "ancre seule", avec le sujet
      d'article à créer, pour un maillage futur.

    Paramètres :
      site_pages : liste de {"url", "label"} issue de fetch_site_urls() (peut être vide).

    Retour :
    {
      "links": [
        {
          "anchor": "texte d'ancre",
          "url": "https://..."  ou  null,
          "type": "existing" | "to_create",
          "target_section": "H2: ... (où placer le lien dans l'article)",
          "reason": "justification courte"
        }
      ]
    }
    """
    if not keyword or not anthropic_key:
        return {"links": []}

    site_pages = site_pages or []
    lang = _lang_label(language_code)
    outline = _structure_to_outline(structure)

    # On limite la liste des pages envoyée au modèle (tokens)
    pages_for_prompt = site_pages[:150]
    if pages_for_prompt:
        pages_block = "\n".join(f"- {p['url']}" for p in pages_for_prompt)
        pages_instr = f"""Voici les pages EXISTANTES du site cible (utilise EXACTEMENT ces URLs
pour les liens de type "existing", sans en inventer d'autres) :
{pages_block}"""
    else:
        pages_instr = ("Aucune page existante n'a pu être récupérée sur le site cible. "
                       "Toutes les recommandations doivent donc être de type \"to_create\" "
                       "(ancre seule, sans URL).")

    prompt = f"""Tu es un expert SEO en maillage interne. Nouvel article à publier.
Mot-clé principal : "{keyword}".

Plan (structure Hn) de l'article :
{outline}

{pages_instr}

Propose jusqu'à {max_links} recommandations de LIENS INTERNES en {lang}, pertinents
sémantiquement avec le sujet. Deux cas :

1. Une page existante ci-dessus est pertinente -> type "existing" :
   fournis "anchor" (ancre optimisée, variée, non sur-optimisée) ET "url" (l'URL exacte
   de la liste).
2. Aucune page existante ne couvre le sujet, mais un lien serait utile -> type "to_create" :
   fournis "anchor" et propose dans "reason" le sujet/article à créer. "url" = null.

Pour chaque reco, indique la section de l'article ("target_section") où insérer le lien.
Évite les ancres génériques ("cliquez ici"). Priorise la pertinence à la quantité.

Réponds UNIQUEMENT avec un objet JSON valide, sans texte ni Markdown, au format :
{{
  "links": [
    {{
      "anchor": "...",
      "url": "https://..." ou null,
      "type": "existing" ou "to_create",
      "target_section": "H2: ...",
      "reason": "..."
    }}
  ]
}}"""

    try:
        raw = _call_claude(prompt, anthropic_key, model, max_tokens=1800)
    except Exception as e:  # noqa: BLE001
        return {"error": f"Erreur API (maillage interne) : {e}", "links": []}

    data = _parse_json(raw)
    if not isinstance(data, dict) or "links" not in data:
        if isinstance(data, list):
            data = {"links": data}
        else:
            return {"error": "Réponse non exploitable pour le maillage interne.",
                    "links": []}

    valid_urls = {p["url"] for p in site_pages}
    links = []
    for l in data.get("links", []):
        if not isinstance(l, dict):
            continue
        anchor = str(l.get("anchor", "")).strip()
        if not anchor:
            continue
        url = l.get("url")
        url = str(url).strip() if url else None
        ltype = str(l.get("type", "")).strip().lower()

        # Garde-fou : un lien "existing" doit pointer vers une URL réellement présente.
        if url and url not in valid_urls:
            # URL hallucinée -> on rétrograde en "à créer" (ancre seule)
            url = None
        if url:
            ltype = "existing"
        else:
            ltype = "to_create"

        links.append({
            "anchor": anchor,
            "url": url,
            "type": ltype,
            "target_section": str(l.get("target_section", "")).strip(),
            "reason": str(l.get("reason", "")).strip(),
        })

    # Tri : liens existants d'abord, puis ancres à créer
    links.sort(key=lambda x: 0 if x["type"] == "existing" else 1)
    return {"links": links[:max_links]}


# ══════════════════════════════════════════════════════════════════════════════
# 2 bis) MAILLAGE INTERNE MANUEL (l'utilisateur définit lui-même ancres + liens)
# ══════════════════════════════════════════════════════════════════════════════
def parse_manual_internal_links(data) -> dict:
    """
    Transforme une saisie manuelle d'ancres/liens en structure exploitable,
    identique à celle renvoyée par generate_internal_linking().

    Accepte :
      - une chaîne de texte, une entrée par ligne, au format :
            ancre | https://url            -> ancre + lien   (type "existing")
            ancre                          -> ancre seule    (type "to_create")
        Séparateurs acceptés entre ancre et URL : « | », tabulation, « ; ».
      - une liste de dicts (ex. depuis un tableau éditable) avec les clés
        "ancre"/"anchor" et "lien"/"url".

    Retour : {"links": [{"anchor", "url"|None, "type", "target_section", "reason"}]}
    """
    rows = []

    if isinstance(data, str):
        for line in data.splitlines():
            line = line.strip()
            if not line:
                continue
            # Cherche le premier séparateur présent
            sep = next((s for s in ("|", "\t", ";") if s in line), None)
            if sep:
                anchor, _, url = line.partition(sep)
                rows.append((anchor.strip(), url.strip()))
            else:
                rows.append((line, ""))
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                anchor = str(item.get("ancre", item.get("anchor", "")) or "").strip()
                url = str(item.get("lien", item.get("url", "")) or "").strip()
                if anchor:
                    rows.append((anchor, url))
            elif isinstance(item, (list, tuple)) and item:
                anchor = str(item[0]).strip()
                url = str(item[1]).strip() if len(item) > 1 and item[1] else ""
                if anchor:
                    rows.append((anchor, url))

    links = []
    for anchor, url in rows:
        if not anchor:
            continue
        # Normalise une éventuelle URL sans schéma
        if url and not re.match(r"^https?://", url, re.I) and "." in url:
            url = "https://" + url
        is_link = bool(url) and re.match(r"^https?://", url or "", re.I) is not None
        links.append({
            "anchor": anchor,
            "url": url if is_link else None,
            "type": "existing" if is_link else "to_create",
            "target_section": "",
            "reason": "",
        })

    links.sort(key=lambda x: 0 if x["type"] == "existing" else 1)
    return {"links": links}


# ══════════════════════════════════════════════════════════════════════════════
# 3) RENDU (HTML pour copier-coller + injection dans le .docx)
# ══════════════════════════════════════════════════════════════════════════════
def enrichment_to_html(secondary_keywords: dict = None, internal_links: dict = None) -> str:
    """Bloc HTML des sections d'enrichissement (pour l'aperçu / copier-coller)."""
    parts = []

    clusters = (secondary_keywords or {}).get("clusters", [])
    if clusters:
        parts.append("<h2>Mots-clés secondaires</h2>")
        for c in clusters:
            theme = _html.escape(c.get("theme", ""))
            placement = _html.escape(c.get("placement", ""))
            kws = ", ".join(_html.escape(k) for k in c.get("keywords", []))
            parts.append(f"<p><strong>{theme}</strong>"
                         + (f" <em>({placement})</em>" if placement else "")
                         + f"<br>{kws}</p>")

    links = (internal_links or {}).get("links", [])
    if links:
        parts.append("<h2>Maillage interne</h2>")
        existing = [l for l in links if l.get("type") == "existing"]
        to_create = [l for l in links if l.get("type") != "existing"]

        if existing:
            parts.append("<h3>Liens vers des pages existantes (ancre + lien)</h3><ul>")
            for l in existing:
                anchor = _html.escape(l.get("anchor", ""))
                url = _html.escape(l.get("url", ""))
                section = _html.escape(l.get("target_section", ""))
                parts.append(
                    f'<li><a href="{url}">{anchor}</a> → {url}'
                    + (f" <em>[{section}]</em>" if section else "")
                    + "</li>"
                )
            parts.append("</ul>")

        if to_create:
            parts.append("<h3>Ancres à créer (contenu à produire — pas encore de lien)</h3><ul>")
            for l in to_create:
                anchor = _html.escape(l.get("anchor", ""))
                section = _html.escape(l.get("target_section", ""))
                reason = _html.escape(l.get("reason", ""))
                parts.append(
                    f"<li><strong>{anchor}</strong>"
                    + (f" <em>[{section}]</em>" if section else "")
                    + (f" — {reason}" if reason else "")
                    + "</li>"
                )
            parts.append("</ul>")

    return "".join(parts)


def append_enrichment_to_docx(docx_buffer, secondary_keywords: dict = None,
                              internal_links: dict = None):
    """
    Ajoute les sections d'enrichissement à un .docx déjà généré.

    docx_buffer : BytesIO renvoyé par export_to_docx().
    Retourne un NOUVEAU BytesIO (positionné au début) contenant le document enrichi.
    Ne modifie pas le module d'export d'origine.
    """
    import io
    from docx import Document

    clusters = (secondary_keywords or {}).get("clusters", [])
    links = (internal_links or {}).get("links", [])
    if not clusters and not links:
        docx_buffer.seek(0)
        return docx_buffer

    docx_buffer.seek(0)
    doc = Document(docx_buffer)

    if clusters:
        doc.add_heading("Mots-clés secondaires", level=1)
        for c in clusters:
            theme = c.get("theme", "")
            placement = c.get("placement", "")
            head = theme + (f" ({placement})" if placement else "")
            doc.add_heading(head, level=2)
            doc.add_paragraph(", ".join(c.get("keywords", [])))

    if links:
        doc.add_heading("Maillage interne", level=1)
        existing = [l for l in links if l.get("type") == "existing"]
        to_create = [l for l in links if l.get("type") != "existing"]

        if existing:
            doc.add_heading("Liens vers des pages existantes (ancre + lien)", level=2)
            for l in existing:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(l.get("anchor", "")).bold = True
                p.add_run(f"  →  {l.get('url', '')}")
                if l.get("target_section"):
                    p.add_run(f"  [{l['target_section']}]").italic = True

        if to_create:
            doc.add_heading("Ancres à créer (contenu à produire — pas encore de lien)", level=2)
            for l in to_create:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(l.get("anchor", "")).bold = True
                extra = []
                if l.get("target_section"):
                    extra.append(l["target_section"])
                if l.get("reason"):
                    extra.append(l["reason"])
                if extra:
                    p.add_run("  — " + " · ".join(extra))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
