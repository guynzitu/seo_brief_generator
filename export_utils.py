"""
Export utilities for SEO briefs.
Supports DOCX and Markdown (Surfer SEO / SEO Quantum formats).
"""
import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def export_to_docx(brief: dict, metadata: dict) -> io.BytesIO:
    """
    Export a brief to a .docx file (compatible Google Docs import).

    Args:
        brief: the generated brief dict
        metadata: dict with keyword, word_range, avg_words, competitor_data

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)

    # ── Title page ────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"Brief SEO")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Mot-clé : {metadata.get('keyword', '')}")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    # Metadata
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"Nombre de mots recommandé : {metadata.get('word_range', 'N/A')} | "
        f"Moyenne concurrents : {metadata.get('avg_words', 'N/A')} mots | "
        f"Pages analysées : {len(metadata.get('competitor_data', []))}"
    )
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    # ── Section: URLs analyzed ────────────────────────────────────────────
    competitor_data = metadata.get("competitor_data", [])
    if competitor_data:
        doc.add_heading("URLs analysées (Top concurrents)", level=1)
        for i, comp in enumerate(competitor_data, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {comp['url']}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            word_hint = p.add_run(f"  — {comp.get('word_count', 0)} mots")
            word_hint.font.size = Pt(9)
            word_hint.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()  # spacer

    # ── Section: SEO Tags ─────────────────────────────────────────────────
    doc.add_heading("Balises SEO", level=1)

    # Title
    doc.add_heading("Title tag", level=2)
    p = doc.add_paragraph()
    run = p.add_run(brief.get("title", "N/A"))
    run.font.bold = True
    run.font.size = Pt(12)
    char_count = len(brief.get("title", ""))
    hint = doc.add_paragraph()
    hint_run = hint.add_run(f"({char_count} caractères — recommandé : 50-70)")
    hint_run.font.size = Pt(9)
    hint_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Meta description
    doc.add_heading("Meta description", level=2)
    p = doc.add_paragraph()
    run = p.add_run(brief.get("meta_description", "N/A"))
    run.font.bold = True
    run.font.size = Pt(12)
    char_count = len(brief.get("meta_description", ""))
    hint = doc.add_paragraph()
    hint_run = hint.add_run(f"({char_count} caractères — recommandé : 140-150)")
    hint_run.font.size = Pt(9)
    hint_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # H1
    doc.add_heading("H1 — Titre de l'article", level=2)
    p = doc.add_paragraph()
    run = p.add_run(brief.get("h1", "N/A"))
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    # ── Section: Content Structure ────────────────────────────────────────
    doc.add_heading("Structure du contenu", level=1)

    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        title = item.get("title", "")
        description = item.get("description", "")

        level_num = int(level[1]) if len(level) == 2 else 2
        # Map to docx heading levels (H2 -> level 2, H3 -> level 3, etc.)
        heading_level = min(level_num, 4)

        h = doc.add_heading(f"{level} : {title}", level=heading_level)

        if description:
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(f"→ {description}")
            desc_run.font.size = Pt(10)
            desc_run.font.italic = True
            desc_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── Section: FAQ ──────────────────────────────────────────────────────
    if brief.get("faq"):
        doc.add_heading("FAQ", level=2)  # H2
        for i, q in enumerate(brief["faq"], 1):
            question = q.get("question", "")
            guideline = q.get("answer_guideline", "")

            doc.add_heading(f"{question}", level=3)  # H3

            if guideline:
                a_para = doc.add_paragraph()
                a_run = a_para.add_run(f"→ {guideline}")
                a_run.font.size = Pt(10)
                a_run.font.italic = True
                a_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── Section: Word count ───────────────────────────────────────────────
    doc.add_heading("Recommandation de longueur", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Nombre de mots recommandé : ").font.size = Pt(11)
    run = p.add_run(f"{metadata.get('word_range', 'N/A')} mots")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p2 = doc.add_paragraph()
    p2.add_run(
        f"Basé sur l'analyse de {len(metadata.get('competitor_data', []))} "
        f"pages concurrentes (moyenne : {metadata.get('avg_words', 'N/A')} mots)."
    ).font.size = Pt(10)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_markdown(brief: dict, metadata: dict, format_type: str = "surfer") -> str:
    """
    Export a brief to Markdown format compatible with SEO tools.

    Args:
        brief: the generated brief dict
        metadata: dict with keyword, word_range, avg_words, competitor_data
        format_type: "surfer" for Surfer SEO format, "seoquantum" for SEO Quantum

    Returns:
        Markdown string
    """
    keyword = metadata.get("keyword", "")
    word_range = metadata.get("word_range", "N/A")
    avg_words = metadata.get("avg_words", "N/A")
    competitor_data = metadata.get("competitor_data", [])
    pages_count = len(competitor_data)

    if format_type == "surfer":
        return _format_surfer(brief, keyword, word_range, avg_words, pages_count, competitor_data)
    else:
        return _format_seoquantum(brief, keyword, word_range, avg_words, pages_count, competitor_data)


def _format_surfer(brief, keyword, word_range, avg_words, pages_count, competitor_data):
    """Surfer SEO compatible markdown format."""
    lines = []
    lines.append(f"# Brief SEO — {keyword}")
    lines.append("")
    lines.append(f"**Mot-clé cible** : {keyword}")
    lines.append(f"**Nombre de mots recommandé** : {word_range}")
    lines.append(f"**Moyenne concurrents** : {avg_words} mots ({pages_count} pages analysées)")
    lines.append("")

    # URLs analyzed
    if competitor_data:
        lines.append("## URLs analysées (Top concurrents)")
        lines.append("")
        for i, comp in enumerate(competitor_data, 1):
            lines.append(f"{i}. {comp['url']} — {comp.get('word_count', 0)} mots")
        lines.append("")

    # SEO Tags
    lines.append("---")
    lines.append("")
    lines.append("## Balises SEO")
    lines.append("")
    lines.append(f"**Title** : {brief.get('title', 'N/A')}")
    lines.append(f"  _{len(brief.get('title', ''))} caractères_")
    lines.append("")
    lines.append(f"**Meta description** : {brief.get('meta_description', 'N/A')}")
    lines.append(f"  _{len(brief.get('meta_description', ''))} caractères_")
    lines.append("")

    # H1
    lines.append("---")
    lines.append("")
    lines.append(f"H1 : {brief.get('h1', 'N/A')}")
    lines.append("")

    # Structure
    lines.append("---")
    lines.append("")
    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        title = item.get("title", "")
        description = item.get("description", "")

        lines.append(f"{level} : {title}")
        if description:
            lines.append(f"  > {description}")
        lines.append("")

    # FAQ
    if brief.get("faq"):
        lines.append("---")
        lines.append("")
        lines.append("## FAQ")
        lines.append("")
        for q in brief["faq"]:
            lines.append(f"**Q : {q.get('question', '')}**")
            if q.get("answer_guideline"):
                lines.append(f"  > Guideline : {q['answer_guideline']}")
            lines.append("")

    # Word count
    lines.append("---")
    lines.append("")
    lines.append(f"**Nombre de mots à rédiger** : {word_range} mots")
    lines.append("")

    return "\n".join(lines)


def _format_seoquantum(brief, keyword, word_range, avg_words, pages_count, competitor_data):
    """SEO Quantum compatible markdown format with explicit H tags."""
    lines = []
    lines.append(f"BRIEF SEO — {keyword}")
    lines.append(f"Mot-clé cible : {keyword}")
    lines.append(f"Nombre de mots : {word_range}")
    lines.append(f"Moyenne concurrents : {avg_words} mots")
    lines.append("")

    # URLs analyzed
    if competitor_data:
        lines.append("== URLS ANALYSEES (TOP CONCURRENTS) ==")
        for i, comp in enumerate(competitor_data, 1):
            lines.append(f"{i}. {comp['url']} — {comp.get('word_count', 0)} mots")
        lines.append("")

    # SEO Tags
    lines.append("== BALISES SEO ==")
    lines.append(f"TITLE : {brief.get('title', 'N/A')} ({len(brief.get('title', ''))} car.)")
    lines.append(f"META DESCRIPTION : {brief.get('meta_description', 'N/A')} ({len(brief.get('meta_description', ''))} car.)")
    lines.append("")

    # Structure with explicit tags
    lines.append("== STRUCTURE DU CONTENU ==")
    lines.append("")
    lines.append(f"H1 : {brief.get('h1', 'N/A')}")
    lines.append("")

    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        title = item.get("title", "")
        description = item.get("description", "")

        lines.append(f"{level} : {title}")
        if description:
            lines.append(f"   Contenu attendu : {description}")
        lines.append("")

    # FAQ
    if brief.get("faq"):
        lines.append("== FAQ ==")
        lines.append("")
        for i, q in enumerate(brief["faq"], 1):
            lines.append(f"Q{i} : {q.get('question', '')}")
            if q.get("answer_guideline"):
                lines.append(f"   Réponse attendue : {q['answer_guideline']}")
            lines.append("")

    # Word count
    lines.append("== LONGUEUR ==")
    lines.append(f"Nombre de mots à rédiger : {word_range} mots")
    lines.append("")

    return "\n".join(lines)


def export_to_plain_text(brief: dict, metadata: dict) -> str:
    """
    Export a brief to plain text format — universal, copy-paste friendly.
    No markdown syntax, no special characters.
    """
    keyword = metadata.get("keyword", "")
    word_range = metadata.get("word_range", "N/A")
    avg_words = metadata.get("avg_words", "N/A")
    competitor_data = metadata.get("competitor_data", [])

    lines = []
    lines.append(f"BRIEF SEO : {keyword}")
    lines.append(f"Nombre de mots recommandé : {word_range}")
    lines.append(f"Moyenne concurrents : {avg_words} mots")
    lines.append("")

    # URLs
    if competitor_data:
        lines.append("URLS ANALYSEES :")
        for i, comp in enumerate(competitor_data, 1):
            lines.append(f"  {i}. {comp['url']} ({comp.get('word_count', 0)} mots)")
        lines.append("")

    # SEO Tags
    lines.append("BALISES SEO")
    lines.append(f"  Title : {brief.get('title', 'N/A')} ({len(brief.get('title', ''))} caractères)")
    lines.append(f"  Meta description : {brief.get('meta_description', 'N/A')} ({len(brief.get('meta_description', ''))} caractères)")
    lines.append("")

    # Structure
    lines.append("STRUCTURE DU CONTENU")
    lines.append(f"  H1 : {brief.get('h1', 'N/A')}")
    lines.append("")

    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        title = item.get("title", "")
        description = item.get("description", "")
        level_num = int(level[1]) if len(level) == 2 else 2
        indent = "  " * (level_num - 1)

        lines.append(f"{indent}{level} : {title}")
        if description:
            lines.append(f"{indent}  > {description}")
    lines.append("")

    # FAQ
    if brief.get("faq"):
        lines.append("FAQ")
        for i, q in enumerate(brief["faq"], 1):
            lines.append(f"  Q{i} : {q.get('question', '')}")
            if q.get("answer_guideline"):
                lines.append(f"    Guideline : {q['answer_guideline']}")
        lines.append("")

    # Word count
    lines.append(f"LONGUEUR : {word_range} mots")

    return "\n".join(lines)


def export_to_html(brief: dict, metadata: dict) -> str:
    """
    Export a brief to rich HTML format.
    When copied and pasted into Google Docs, Word, Notion, etc.,
    the H1/H2/H3 tags are preserved as real formatted headings.
    """
    keyword = metadata.get("keyword", "")
    word_range = metadata.get("word_range", "N/A")
    avg_words = metadata.get("avg_words", "N/A")
    competitor_data = metadata.get("competitor_data", [])

    parts = []

    # Header info
    parts.append(f'<p style="color:#64748b;font-size:13px;margin-bottom:4px;"><strong>Mot-clé cible :</strong> {keyword}</p>')
    parts.append(f'<p style="color:#64748b;font-size:13px;margin-bottom:4px;"><strong>Nombre de mots recommandé :</strong> {word_range}</p>')
    parts.append(f'<p style="color:#64748b;font-size:13px;margin-bottom:16px;"><strong>Moyenne concurrents :</strong> {avg_words} mots ({len(competitor_data)} pages)</p>')

    # URLs analyzed
    if competitor_data:
        parts.append('<p style="font-size:14px;font-weight:700;color:#334155;margin-bottom:8px;">URLs analysées :</p>')
        parts.append('<ol style="margin:0 0 16px 20px;padding:0;font-size:13px;color:#475569;">')
        for comp in competitor_data:
            parts.append(f'<li>{comp["url"]} — {comp.get("word_count", 0)} mots</li>')
        parts.append('</ol>')

    parts.append('<hr style="border:none;border-top:2px solid #e2e8f0;margin:16px 0;">')

    # SEO Tags
    parts.append('<p style="font-size:15px;font-weight:700;color:#0f172a;margin-bottom:8px;">BALISES SEO</p>')
    title_text = brief.get("title", "N/A")
    meta_text = brief.get("meta_description", "N/A")
    parts.append(f'<p style="margin:4px 0;"><strong>Title :</strong> {title_text} <span style="color:#94a3b8;font-size:12px;">({len(title_text)} car.)</span></p>')
    parts.append(f'<p style="margin:4px 0 16px 0;"><strong>Meta description :</strong> {meta_text} <span style="color:#94a3b8;font-size:12px;">({len(meta_text)} car.)</span></p>')

    parts.append('<hr style="border:none;border-top:2px solid #e2e8f0;margin:16px 0;">')

    # H1
    h1_text = brief.get("h1", "N/A")
    parts.append(f'<h1 style="font-size:24px;font-weight:700;color:#0f172a;margin:16px 0 12px 0;">H1 : {h1_text}</h1>')

    # Structure Hn
    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        title = item.get("title", "")
        description = item.get("description", "")
        level_num = int(level[1]) if len(level) == 2 else 2

        if level_num == 2:
            parts.append(f'<h2 style="font-size:20px;font-weight:700;color:#1e293b;margin:20px 0 6px 0;">H2 : {title}</h2>')
        elif level_num == 3:
            parts.append(f'<h3 style="font-size:17px;font-weight:600;color:#334155;margin:14px 0 4px 16px;">H3 : {title}</h3>')
        elif level_num == 4:
            parts.append(f'<h4 style="font-size:15px;font-weight:600;color:#475569;margin:10px 0 4px 32px;">H4 : {title}</h4>')
        else:
            parts.append(f'<p style="font-weight:600;color:#475569;margin:8px 0 4px {16*(level_num-2)}px;">{level} : {title}</p>')

        if description:
            indent = 16 * (level_num - 1)
            parts.append(f'<p style="color:#64748b;font-size:13px;font-style:italic;margin:0 0 8px {indent}px;">→ {description}</p>')

    # FAQ
    if brief.get("faq"):
        parts.append('<hr style="border:none;border-top:2px solid #e2e8f0;margin:20px 0;">')
        parts.append('<h2 style="font-size:20px;font-weight:700;color:#1e293b;margin:16px 0 12px 0;">H2 : FAQ</h2>')
        for i, q in enumerate(brief["faq"], 1):
            question = q.get("question", "")
            guideline = q.get("answer_guideline", "")
            parts.append(f'<h3 style="font-size:16px;font-weight:600;color:#334155;margin:12px 0 4px 0;">H3 : {question}</h3>')
            if guideline:
                parts.append(f'<p style="color:#64748b;font-size:13px;font-style:italic;margin:0 0 10px 16px;">Guideline : {guideline}</p>')

    # Word count
    parts.append('<hr style="border:none;border-top:2px solid #e2e8f0;margin:16px 0;">')
    parts.append(f'<p style="font-size:14px;"><strong>Nombre de mots à rédiger :</strong> <span style="color:#0f766e;font-weight:700;">{word_range} mots</span></p>')

    return "\n".join(parts)


def export_to_raw_html(brief: dict, metadata: dict) -> str:
    """
    Export brief as clean semantic HTML code (no inline styles).
    For copy-pasting raw HTML into CMS, SEO tools, or any HTML editor.
    """
    keyword = metadata.get("keyword", "")
    word_range = metadata.get("word_range", "N/A")
    avg_words = metadata.get("avg_words", "N/A")
    competitor_data = metadata.get("competitor_data", [])

    lines = []

    # SEO Tags
    title_text = brief.get("title", "")
    meta_text = brief.get("meta_description", "")
    lines.append(f'<p><strong>Title :</strong> {title_text} ({len(title_text)} car.)</p>')
    lines.append(f'<p><strong>Meta description :</strong> {meta_text} ({len(meta_text)} car.)</p>')
    lines.append(f'<p><strong>Nombre de mots recommandé :</strong> {word_range} mots</p>')

    # URLs analyzed
    if competitor_data:
        lines.append('<p><strong>URLs analysées :</strong></p>')
        lines.append("<ul>")
        for comp in competitor_data:
            lines.append(f'  <li><a href="{comp["url"]}">{comp["url"]}</a> — {comp.get("word_count", 0)} mots</li>')
        lines.append("</ul>")

    lines.append('<p><strong>______________________________________________________________</strong></p>')

    # H1
    h1_text = brief.get("h1", "")
    lines.append(f'<h1>H1 : {h1_text}</h1>')

    # Structure
    for item in brief.get("structure", []):
        level = item.get("level", "H2")
        level_lower = level.lower()
        title = item.get("title", "")
        description = item.get("description", "")

        lines.append(f"<{level_lower}>{level} : {title}</{level_lower}>")
        if description:
            lines.append(f"<p><em>→ {description}</em></p>")

    # FAQ (H2 + H3)
    if brief.get("faq"):
        lines.append("<h2>H2 : FAQ</h2>")
        for q in brief["faq"]:
            question = q.get("question", "")
            guideline = q.get("answer_guideline", "")
            lines.append(f"<h3>H3 : {question}</h3>")
            if guideline:
                lines.append(f"<p><em>Guideline : {guideline}</em></p>")

    return "\n".join(lines)
